# -*- coding: utf-8 -*-
"""Письмо бюро: замечания подачи одним документом.

Отдел экспертизы отправляет бюро не таблицу, а письмо: объект, подача,
перечень замечаний по томам, подпись. Здесь оно и собирается — из тех
самых замечаний, что эксперт завёл на вкладках паспорта и сверки.

Формулировки берутся как есть: их писал (или правил) человек. Машина
отвечает только за порядок, нумерацию и шапку.
"""
import io
from datetime import datetime

from docx import Document as Docx
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FONT = 'Times New Roman'
SIZE = Pt(11)

LEVEL_WORDS = {'red': 'расхождение', 'amber': 'требует уточнения',
               'ok': 'замечание'}

MONTHS = ('января февраля марта апреля мая июня июля августа сентября '
          'октября ноября декабря').split()

# колонки таблицы замечаний: номер, предмет, формулировка, листы
WIDTHS = (Cm(1.2), Cm(5.6), Cm(15.9), Cm(2.6))


def _date(dt=None):
    dt = dt or datetime.now()
    return f'{dt.day} {MONTHS[dt.month - 1]} {dt.year} г.'


def _style(doc):
    """Обычный текст письма: Times New Roman 11, без интервалов между абзацами.

    Кириллице нужен явный eastAsia-шрифт, иначе Word подставляет свой
    и документ у получателя выглядит не так, как у отправителя.
    """
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = SIZE
    rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rpr.set(qn(attr), FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def _para(doc, text='', bold=False, align=None, size=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p


def _fixed(table, widths):
    """Фиксированная раскладка колонок.

    Одних cell.width мало: пока таблица в режиме autofit, Word и
    LibreOffice раскладывают колонки по содержимому, и колонка с
    формулировкой оказывается уже колонки «№».
    """
    table.autofit = False
    tbl = table._tbl
    tbl.tblPr.append(_el('w:tblLayout', {'w:type': 'fixed'}))
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl.remove(grid)
    grid = _el('w:tblGrid')
    for w in widths:
        grid.append(_el('w:gridCol', {'w:w': str(int(w.twips))}))
    tbl.insert(1, grid)
    return table


def _el(tag, attrs=None):
    from docx.oxml import OxmlElement
    el = OxmlElement(tag)
    for k, v in (attrs or {}).items():
        el.set(qn(k), v)
    return el


def _cell(cell, text, bold=False, size=Pt(10), align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text or '')
    run.bold = bold
    run.font.size = size
    run.font.name = FONT
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    return cell


def build(project, submission, groups, author='', org_name='',
          total=0) -> bytes:
    """groups: [(том, [замечания])] в порядке вывода. -> байты .docx

    Лист альбомный: в таблице замечаний живёт формулировка на две-три
    строки, и на книжной странице она превращается в лапшу.
    """
    doc = Docx()
    _style(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.8)
    section.top_margin = section.bottom_margin = Cm(1.5)

    _para(doc, org_name or 'Отдел внутренней экспертизы',
          align=WD_ALIGN_PARAGRAPH.RIGHT, size=Pt(10), space_after=Pt(0))
    _para(doc, _date(), align=WD_ALIGN_PARAGRAPH.RIGHT, size=Pt(10))
    if project.bureau:
        _para(doc, f'В {project.bureau}', bold=True, space_after=Pt(0))
    _para(doc, 'Замечания по результатам внутренней экспертизы '
               'рабочей документации', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    head = [f'Объект: {project.name}']
    if project.code:
        head.append(f'Шифр: {project.code}')
    if submission is not None:
        head.append(f'Подача: {submission.label}')
    for line in head:
        _para(doc, line, space_after=Pt(0))
    _para(doc, f'Всего замечаний: {total}')

    n = 0
    for document, remarks in groups:
        if not remarks:
            continue
        title = document.cipher or document.filename
        section_title = f'{title} — {document.section_label or document.section}'
        _para(doc, section_title.rstrip(' —'), bold=True)

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fixed(table, WIDTHS)
        for cell, w in zip(table.rows[0].cells, WIDTHS):
            cell.width = w
        hdr = table.rows[0].cells
        for cell, text in zip(hdr, ('№', 'Позиция / проверка', 'Замечание',
                                    'Листы')):
            _cell(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        for r in remarks:
            n += 1
            row = table.add_row()
            cells = row.cells
            for cell, w in zip(cells, WIDTHS):
                cell.width = w
            _cell(cells[0], str(n), align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(cells[1], r.subject)
            _cell(cells[2], r.text)
            _cell(cells[3], ', '.join(str(p) for p in (r.sheets or [])),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(doc, '')

    _para(doc, 'Замечания составлены по результатам проверки состава '
               'документации и сверки спецификаций с чертежами. Просим '
               'учесть их при выпуске следующего изменения.')
    _para(doc, '')
    p = _para(doc, 'Эксперт отдела внутренней экспертизы')
    p.add_run('\t\t\t')
    p.add_run(f'________________ / {author or ""}')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def filename(project, submission):
    code = project.code or project.name
    label = submission.label if submission is not None else ''
    return f'Замечания_{code}_{label}.docx'.replace(' ', '_')
